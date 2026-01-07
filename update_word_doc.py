"""
Script to update the Equity Ranking With Ridge Stacking.docx with latest results
"""
from docx import Document
import re

def update_document():
    doc_path = r"D:\trade\Equity Ranking With Ridge Stacking.docx"
    doc = Document(doc_path)

    # Track changes made
    changes_made = []

    # 1. Update Abstract - Find and replace the XGBoost paragraph
    old_abstract_text = "In our feature combination pipeline, XGBoost consistently outperforms other base learners, achieving an average top-decile forward return of approximately 6.76% per 10-day period in the most optimized configuration (Combo 0060), compared to 2.21% for linear baselines"

    new_abstract_text = "In our feature combination pipeline, LambdaRank consistently outperforms other base learners after correcting for time leakage, achieving an average top-decile forward return of approximately 4.30% per 10-day period. This is a significant improvement over the 0.79% achieved by the linear baseline (Elastic Net). The results confirm that while non-linear stacking extracts substantial alpha, ranking-specific objectives (LambdaRank) provide superior robustness compared to point-wise regression learners (XGBoost) when strict temporal purging and embargoing protocols are applied."

    # Search through all paragraphs
    for para in doc.paragraphs:
        if "XGBoost consistently outperforms other base learners" in para.text:
            # Replace the text
            for run in para.runs:
                if "XGBoost consistently outperforms" in run.text:
                    run.text = run.text.replace(old_abstract_text, new_abstract_text)
                    changes_made.append("Updated Abstract paragraph")
                    break

    # 2. Find and update tables (Table 1: Performance Summary)
    for table in doc.tables:
        # Check if this is the performance table by looking for column headers
        if len(table.rows) > 0:
            header_row = table.rows[0]
            header_text = ' '.join([cell.text for cell in header_row.cells])

            if "Model" in header_text and "Avg Top Return" in header_text:
                # This is our performance table - update it
                # Clear existing data rows (keep header)
                while len(table.rows) > 1:
                    table._element.remove(table.rows[-1]._element)

                # Add updated data rows
                performance_data = [
                    ("LambdaRank", "4.30%", "4.14%", "188.7%", "0.0176", "1.99"),
                    ("Ridge Stacking", "3.04%", "2.85%", "112.5%", "-0.0152", "1.66"),
                    ("XGBoost", "2.75%", "2.58%", "98.2%", "-0.0107", "1.49"),
                    ("CatBoost", "1.19%", "1.01%", "34.9%", "-0.0442", "1.31"),
                    ("Elastic Net", "0.79%", "0.60%", "21.8%", "-0.0348", "0.85"),
                ]

                for row_data in performance_data:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row_data):
                        if i < len(row_cells):
                            row_cells[i].text = value
                            # Bold the LambdaRank row
                            if row_data[0] == "LambdaRank":
                                for run in row_cells[i].paragraphs[0].runs:
                                    run.bold = True

                changes_made.append("Updated Performance Summary Table")
                break

    # 3. Add/Update Analysis & Discussion section
    analysis_text = """The "Leakage Decay" Effect:
The implementation of a strict temporal fix revealed that regression-based models (XGBoost and CatBoost) were significantly benefiting from look-ahead bias in the uncorrected dataset. XGBoost's performance decayed from an initial 6.76% to 2.75% once the target alignment was corrected. In contrast, LambdaRank proved to be the most resilient architecture. By optimizing for the relative cross-sectional rank rather than absolute return values, LambdaRank filters out much of the noise and bias that regression models tend to overfit.

Economic Significance:
Despite the moderation in gross returns, the strategy remains highly viable. With a dynamic Garman-Klass cost model estimating a drag of approximately 16-19 bps per 10-day window, the LambdaRank model retains a net alpha of 4.14%. The Ridge Stacking meta-learner (3.04%) also demonstrates the value of regularized ensemble methods, effectively "smoothing" the signals from the more volatile base regressors."""

    # Find section containing "Analysis" or "Discussion" and update/add content
    found_analysis = False
    for i, para in enumerate(doc.paragraphs):
        if "Analysis" in para.text or "Discussion" in para.text:
            # Check if this is a heading
            if para.style.name.startswith('Heading'):
                # Add analysis text after this heading
                # First, check if next paragraph contains old analysis
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    if "Leakage" in next_para.text or "implementation of a strict temporal fix" in next_para.text:
                        # Replace existing text
                        next_para.text = analysis_text
                        found_analysis = True
                        changes_made.append("Updated Analysis & Discussion section")
                        break

    if not found_analysis:
        # Add new section at the end
        doc.add_heading("Analysis & Discussion", level=2)
        doc.add_paragraph(analysis_text)
        changes_made.append("Added Analysis & Discussion section")

    # 4. Update Key Findings section
    key_findings_text = """• Ranking Robustness: Pairwise ranking objectives (LambdaRank) are superior to point-wise regression (XGBoost) for T+10 horizons, showing much higher resilience to temporal leakage.
• Meta-Learner Value: The Ridge Stacker successfully aggregates multiple signals to outperform individual GBDT regressors, providing a more stable Sharpe ratio (1.66).
• Linear Baseline Failure: Once leakage is removed, the linear Elastic Net baseline (0.79% gross) becomes marginal, highlighting the necessity of non-linear feature interaction in modern equity markets."""

    found_findings = False
    for i, para in enumerate(doc.paragraphs):
        if "Key Findings" in para.text or "Conclusion" in para.text:
            if para.style.name.startswith('Heading'):
                # Update bullet points that follow
                j = i + 1
                # Remove old bullet points
                bullets_to_remove = []
                while j < len(doc.paragraphs) and (doc.paragraphs[j].text.strip().startswith('•') or doc.paragraphs[j].text.strip().startswith('-')):
                    bullets_to_remove.append(j)
                    j += 1

                # Remove old bullets in reverse order
                for idx in reversed(bullets_to_remove):
                    p = doc.paragraphs[idx]._element
                    p.getparent().remove(p)

                # Add new findings after the heading
                for line in key_findings_text.split('\n'):
                    if line.strip():
                        # Insert after heading
                        new_para = doc.add_paragraph(line)
                        # Move to correct position (after heading)
                        heading_element = doc.paragraphs[i]._element
                        heading_element.addnext(new_para._element)

                found_findings = True
                changes_made.append("Updated Key Findings section")
                break

    if not found_findings:
        # Add new section
        doc.add_heading("Key Findings", level=2)
        for line in key_findings_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        changes_made.append("Added Key Findings section")

    # Save the document
    output_path = r"D:\trade\Equity Ranking With Ridge Stacking_updated.docx"
    doc.save(output_path)

    return output_path, changes_made

if __name__ == "__main__":
    try:
        output_path, changes = update_document()
        print(f"Document updated successfully!")
        print(f"Saved to: {output_path}")
        print(f"\nChanges made:")
        for change in changes:
            print(f"  - {change}")
    except Exception as e:
        print(f"Error updating document: {e}")
        import traceback
        traceback.print_exc()
