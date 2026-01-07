#!/usr/bin/env python3
"""
Update an existing paper DOCX by appending reviewer-requested academic sections and embedding
the latest long-only backtest artifacts.

This script does NOT delete/overwrite content; it appends a "Revision Addendum" at the end.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

import pandas as pd


def _read_csv(p: Path) -> pd.DataFrame:
    return pd.read_csv(p)


def _add_heading(doc, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def _add_paragraph(doc, text: str):
    return doc.add_paragraph(text)


def _add_table(doc, df: pd.DataFrame, max_rows: int = 40):
    d = df.copy()
    if len(d) > max_rows:
        d = d.head(max_rows)
    table = doc.add_table(rows=1, cols=len(d.columns))
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    hdr = table.rows[0].cells
    for i, c in enumerate(d.columns):
        hdr[i].text = str(c)
    for _, row in d.iterrows():
        cells = table.add_row().cells
        for i, v in enumerate(row.tolist()):
            if isinstance(v, float):
                cells[i].text = f"{v:.6g}"
            else:
                cells[i].text = str(v)
    return table


def _add_picture_if_exists(doc, path: Path, width_inches: float = 6.5):
    if not path.exists():
        return False
    from docx.shared import Inches

    doc.add_picture(str(path), width=Inches(width_inches))
    return True


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input-docx", required=True)
    ap.add_argument("--output-docx", required=True)
    ap.add_argument("--artifacts-dir", required=True, help="Folder produced by scripts/generate_paper_revision_artifacts.py")
    ap.add_argument("--run-dir", required=True, help="Backtest run folder containing per_model_topN_vs_benchmark*.png/csv")
    ap.add_argument("--prior-performance-report", default=None, help="Optional previous performance_report_*.csv (baseline) to show BEFORE/AFTER comparison.")
    args = ap.parse_args()

    from docx import Document

    in_doc = Path(args.input_docx)
    out_doc = Path(args.output_docx)
    artifacts_dir = Path(args.artifacts_dir)
    run_dir = Path(args.run_dir)

    doc = Document(str(in_doc))

    _add_heading(doc, "Revision Addendum (Reviewer-Requested Deepening)", level=1)
    _add_paragraph(
        doc,
        "This addendum addresses reviewer feedback on theoretical formalization, interpretability, risk attribution, microstructure assumptions, and robustness. "
        "All results below are computed on the time-split TEST window (last 20%) under a long-only Top-N strategy net of transaction costs.",
    )

    # A) Theoretical Formalization
    _add_heading(doc, "A) Theoretical Formalization", level=2)
    _add_heading(doc, "A1) LambdaRank objective (LambdaLoss) vs pointwise MSE", level=3)
    _add_paragraph(
        doc,
        "We optimize a ranking objective by learning pairwise preferences within each date’s cross-section. "
        "For a query date t, items i,j with labels y_i,y_j, and score s_i=f(x_i), define a pairwise probability via a logistic link:",
    )
    _add_paragraph(doc, r"$P_{ij} = \sigma(s_i - s_j)$, where $\sigma(z)=1/(1+e^{-z})$.")
    _add_paragraph(
        doc,
        "A common pairwise loss is logistic cross-entropy over ordered pairs (y_i > y_j):",
    )
    _add_paragraph(doc, r"$\mathcal{L}_{pair} = \sum_{(i,j):y_i>y_j} \log(1+\exp(-(s_i-s_j))).$")
    _add_paragraph(
        doc,
        "LambdaRank modifies gradients (the “lambdas”) to weight pairs by their impact on NDCG@K, "
        "aligning optimization with top-K selection rather than minimizing prediction error magnitude as in pointwise MSE.",
    )

    _add_heading(doc, "A2) Ridge Stacking as meta-learner with L2 regularization", level=3)
    _add_paragraph(
        doc,
        "Let base learners produce cross-sectional scores \\hat{s}^{(1)},\\hat{s}^{(2)},...,\\hat{s}^{(m)} for each (t,i). "
        "Ridge stacking fits a linear meta-learner:",
    )
    _add_paragraph(doc, r"$\hat{y} = \beta_0 + \sum_{k=1}^m \beta_k \hat{s}^{(k)}$")
    _add_paragraph(
        doc,
        "with an L2 penalty to mitigate multicollinearity among base predictions:",
    )
    _add_paragraph(doc, r"$\min_{\beta}\ \sum_n (y_n-\beta_0-\beta^\top \hat{s}_n)^2 + \alpha \|\beta\|_2^2.$")

    # Ridge weights
    ridge_w_csv = artifacts_dir / "ridge_meta_weights.csv"
    ridge_w_png = artifacts_dir / "ridge_meta_weights_top20.png"
    if ridge_w_csv.exists():
        _add_heading(doc, "A3) Ridge meta-learner weight decomposition", level=3)
        _add_paragraph(doc, f"Source: {ridge_w_csv.as_posix()}")
        w = _read_csv(ridge_w_csv)
        _add_table(doc, w, max_rows=10)
        _add_paragraph(doc, "Figure: Ridge meta-learner weights (Top by |weight|).")
        _add_picture_if_exists(doc, ridge_w_png, width_inches=6.0)

    # Make the change obvious: BEFORE vs AFTER (ridge only)
    if args.prior_performance_report:
        prior_path = Path(args.prior_performance_report)
        if prior_path.exists() and perf_csv.exists():
            _add_heading(doc, "A4) Change log: Ridge stacking BEFORE vs AFTER adding LambdaRank", level=3)
            _add_paragraph(doc, f"Prior report (baseline): {prior_path.as_posix()}")
            _add_paragraph(doc, f"New report: {perf_csv.as_posix()}")

            prior = _read_csv(prior_path)
            newrep = _read_csv(perf_csv)
            if "Model" not in prior.columns and "model" in prior.columns:
                prior = prior.rename(columns={"model": "Model"})
            if "Model" not in newrep.columns and "model" in newrep.columns:
                newrep = newrep.rename(columns={"model": "Model"})

            def _row(df: pd.DataFrame, name: str) -> pd.Series:
                rr = df.loc[df["Model"].astype(str) == name].head(1)
                return rr.iloc[0] if len(rr) else pd.Series(dtype=float)

            r0 = _row(prior, "ridge_stacking")
            r1 = _row(newrep, "ridge_stacking")

            metrics = [
                ("avg_top_return_net", "Top-N avg return net (per period)"),
                ("top_sharpe_net", "Top-N Sharpe net (annualized)"),
                ("IC", "IC (Pearson)"),
                ("Rank_IC", "Rank IC (Spearman)"),
                ("avg_top_turnover", "Avg turnover (Top-N)"),
                ("avg_top_cost", "Avg cost (Top-N)"),
            ]

            rows = []
            for k, label in metrics:
                if k in prior.columns and k in newrep.columns:
                    v0 = float(r0.get(k)) if len(r0) else float("nan")
                    v1 = float(r1.get(k)) if len(r1) else float("nan")
                    rows.append(
                        {
                            "metric": k,
                            "description": label,
                            "before": v0,
                            "after": v1,
                            "delta(after-before)": (v1 - v0) if (pd.notna(v0) and pd.notna(v1)) else float("nan"),
                        }
                    )
            if rows:
                comp = pd.DataFrame(rows)
                _add_table(doc, comp, max_rows=40)

    # B) Feature transparency
    _add_heading(doc, "B) Feature Taxonomy & Transparency", level=2)
    feat_csv = artifacts_dir / "feature_list_best_per_model.csv"
    if feat_csv.exists():
        _add_heading(doc, "B1) Core feature list (per model)", level=3)
        feat = _read_csv(feat_csv)
        _add_paragraph(doc, f"Source: {feat_csv.as_posix()}")
        # show top 15 per model for readability
        feat_small = feat.loc[feat["rank"] <= 15].copy()
        _add_table(doc, feat_small, max_rows=80)
    else:
        _add_paragraph(doc, "Feature list artifact NOT FOUND.")

    # C) Long-only performance + risk stats
    _add_heading(doc, "C) Long-only performance, distribution risk, and robustness", level=2)
    perf_csv = artifacts_dir / "performance_report.csv"
    if perf_csv.exists():
        perf = _read_csv(perf_csv)
        cols = [
            "Model",
            "avg_top_return_net",
            "top_sharpe_net",
            "top_win_rate_net",
            "avg_top_turnover",
            "avg_top_cost",
            "IC",
            "Rank_IC",
        ]
        cols = [c for c in cols if c in perf.columns]
        _add_heading(doc, "C1) Summary table (net-of-cost, long-only Top-N)", level=3)
        _add_paragraph(doc, f"Source: {perf_csv.as_posix()}")
        _add_table(doc, perf[cols].copy(), max_rows=10)

    dist_all = artifacts_dir / "dist_stats_all_models.csv"
    if dist_all.exists():
        _add_heading(doc, "C2) Return distribution (skew/kurtosis) — fat-tail diagnostics", level=3)
        dist = _read_csv(dist_all)
        _add_paragraph(doc, f"Source: {dist_all.as_posix()}")
        _add_table(doc, dist, max_rows=20)

    yr_all = artifacts_dir / "yearly_stats_all_models.csv"
    if yr_all.exists():
        _add_heading(doc, "C3) Year-by-year performance (non-overlapping periods)", level=3)
        yr = _read_csv(yr_all)
        # focus on recent years + major regimes
        yr = yr.loc[yr["year"].isin([2020, 2022, 2024, 2025])].copy() if "year" in yr.columns else yr
        _add_paragraph(doc, f"Source: {yr_all.as_posix()}")
        _add_table(doc, yr, max_rows=80)

    # Benchmark plots
    _add_heading(doc, "C4) Benchmark comparison (QQQ) — per-period and cumulative", level=3)
    _add_paragraph(doc, f"Source folder: {run_dir.as_posix()}")
    _add_picture_if_exists(doc, run_dir / "per_model_topN_vs_benchmark.png", width_inches=6.5)
    _add_picture_if_exists(doc, run_dir / "per_model_topN_vs_benchmark_cum.png", width_inches=6.5)

    # D) Microstructure & implementation notes
    _add_heading(doc, "D) Microstructure & Implementation Notes", level=2)
    _add_paragraph(
        doc,
        "Transaction costs are modeled as turnover × cost_bps/1e4 on the long Top-N basket (equal weight). "
        "Turnover is computed from the change in target weights between consecutive rebalance dates.",
    )
    _add_paragraph(
        doc,
        "Capacity and impact are discussed via a square-root market impact approximation: "
        "Impact ≈ Y · σ · sqrt(Q/ADV), with explicit assumptions required for σ, ADV, and participation rate.",
    )

    # E) Limitations / future work
    _add_heading(doc, "E) Limitations and recommended robustness extensions", level=2)
    _add_paragraph(
        doc,
        "IC decay T+1..T+10 requires point-in-time targets for each horizon; current factor dataset exposes T+10 only. "
        "A full hyperparameter sensitivity grid (Ridge α, LambdaRank depth/leaves) is recommended for publication-grade robustness.",
    )

    doc.save(str(out_doc))
    print(f"WROTE {out_doc}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())


