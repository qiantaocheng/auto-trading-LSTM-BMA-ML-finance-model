import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches

BASE_DIR = Path('D:/trade')
MD_PATH = BASE_DIR / 'docs' / 'US_T10_TOPK_RESEARCH_PROCESS.md'
DOCX_PATH = BASE_DIR / 'docs' / 'US_T10_TOPK_RESEARCH_PROCESS.docx'
PERFORMANCE_PATH = BASE_DIR / 'results' / 'paper_costs_nocat_20260106_135011' / 'performance_report_20260106_135130.csv'
BUCKETS_PATH = BASE_DIR / 'results' / 'paper_costs_nocat_20260106_135011' / 'buckets_vs_nasdaq.csv'
OOS_PATH = BASE_DIR / 'results' / 't10_time_split_80_20_costs_20260106_212931' / 'run_20260106_212934' / 'oos_metrics.csv'
TOP20_IMG = BASE_DIR / 'results' / 't10_time_split_80_20_costs_20260106_212931' / 'run_20260106_212934' / 'top20_vs_qqq.png'


def load_markdown_text(path: Path) -> str:
    return path.read_text(encoding='utf-8')


def add_markdown(document: Document, text: str) -> None:
    lines = text.splitlines()
    in_code_block = False
    code_buffer = []
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith('```'):
            if in_code_block:
                paragraph = document.add_paragraph('\n'.join(code_buffer))
                paragraph.style = 'Intense Quote'
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_buffer.append(raw_line)
            continue
        if not stripped:
            document.add_paragraph('')
            continue
        heading_match = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            document.add_heading(text, level=min(level, 4))
            continue
        if stripped == '---':
            document.add_paragraph('')
            continue
        if re.match(r'^-\s+', stripped):
            document.add_paragraph(stripped[2:], style='List Bullet')
            continue
        if re.match(r'^\d+\)', stripped):
            document.add_paragraph(stripped, style='List Number')
            continue
        document.add_paragraph(raw_line)


def fmt_value(value):
    if value is None:
        return ''
    if isinstance(value, float):
        return f"{value:.6g}"
    return str(value)


def add_dataframe_table(document: Document, df: pd.DataFrame, title: str, subtitle: str | None = None) -> None:
    document.add_heading(title, level=3)
    if subtitle:
        document.add_paragraph(subtitle)
    table = document.add_table(rows=1, cols=len(df.columns))
    header_cells = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        header_cells[idx].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            cells[idx].text = fmt_value(row[col])


def build_document() -> None:
    document = Document()
    md_text = load_markdown_text(MD_PATH)
    add_markdown(document, md_text)

    document.add_page_break()
    document.add_heading('Data Appendices', level=2)

    performance_df = pd.read_csv(PERFORMANCE_PATH)
    perf_columns = [
        'Model',
        'N_Predictions',
        'avg_top_return',
        'avg_top_return_net',
        'Rank_IC',
        'IC',
        'long_short_sharpe',
        'win_rate',
        'avg_top_turnover',
        'avg_top_cost',
        'cost_bps'
    ]
    add_dataframe_table(
        document,
        performance_df[perf_columns],
        'Model-level Metrics',
        f'Source: {PERFORMANCE_PATH.as_posix()}'
    )

    buckets_df = pd.read_csv(BUCKETS_PATH)
    latest_bucket = buckets_df.tail(1)[[
        'date',
        'cum_top_1_10_return_net',
        'cum_top_11_20_return_net',
        'cum_top_21_30_return_net',
        'cum_benchmark_return'
    ]]
    add_dataframe_table(
        document,
        latest_bucket,
        'Latest Cumulative Bucket Returns vs Benchmark',
        f'Source: {BUCKETS_PATH.as_posix()} (last row)'
    )

    oos_df = pd.read_csv(OOS_PATH)
    add_dataframe_table(
        document,
        oos_df[[
            'model',
            'top_n',
            'horizon_days',
            'split',
            'test_start',
            'test_end',
            'n_test_rebalances',
            'avg_top_return_pct',
            'avg_top_return_net_pct',
            'avg_benchmark_return_pct',
            'end_cum_top_return_pct',
            'end_cum_top_return_net_pct',
            'end_cum_benchmark_return_pct',
            'avg_top_turnover',
            'avg_top_cost',
            'win_rate',
            'long_short_sharpe'
        ]],
        'Out-of-Sample Top-20 Metrics',
        f'Source: {OOS_PATH.as_posix()}'
    )

    if TOP20_IMG.exists():
        document.add_heading('Result Figure: Top-20 vs QQQ', level=3)
        document.add_picture(str(TOP20_IMG), width=Inches(6.5))
        document.add_paragraph(f'Source: {TOP20_IMG.as_posix()}')
    else:
        document.add_paragraph('Result image not found at: ' + TOP20_IMG.as_posix())

    document.save(DOCX_PATH)


if __name__ == '__main__':
    build_document()
